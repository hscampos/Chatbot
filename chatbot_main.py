import openai
import streamlit as st
from streamlit_chat import message as msg
import docx
import io

with open("api_key.txt", "r") as api_key:
    openai.api_key = api_key.read()

st.title("Chat Bot com Chat GPT")
st.write("***")

if 'conversa' not in st.session_state:
    st.session_state.conversa = []

pergunta = st.text_input("O que deseja perguntar ?")
botaoEnviar = st.button("Perguntar")

# Consulta a OpenAI e trata o retorno.

if botaoEnviar:
    st.session_state.conversa.append({"role": "user", "content": pergunta})
    retornoOpenAi = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=st.session_state.conversa
    )
    st.session_state.conversa.append(
        {"role": "assistant",
         "content": retornoOpenAi['choices'][0]['message']['content']})

# Realiza o construção da conversa com o retorno já tratado e a pergunta feita.

if len(st.session_state.conversa) > 0:
    for i in range(len(st.session_state.conversa)):
        if i % 2 == 0:
            msg("Você: " + st.session_state.conversa[i]['content'], is_user=True)
        else:
            msg("Resposta IA: " + st.session_state.conversa[i]['content'])

# Permite salvar o conteúdo da 'conversa', através de um botão que aparece apenas quando
# o bot efetua uma resposta.

if len(st.session_state.conversa) > 0:
    botaoSalvar = st.button("Salvar Conteúdo")
    if botaoSalvar:
        trabalho = io.BytesIO()
        documento = docx.Document()
        documento.add_heading('Conteúdo Gerado', level=1)

        # Percorre a lista 'conversa' e adiciona ao documento os itens dela,
        # onde a pergunta é um item e a resposta outro.

        for i in range(len(st.session_state.conversa)):
            if i % 2 == 0:
                documento.add_heading("Pergunta", level=2)
                documento.add_paragraph(st.session_state.conversa[i]['content'])
            else:
                documento.add_heading("Resposta", level=2)
                documento.add_paragraph(st.session_state.conversa[i]['content'])

        documento.save(trabalho)
        st.download_button(label="Clique aqui para salvar este conteúdo",
                           data=trabalho,
                           file_name="",
                           mime="application/msword")
